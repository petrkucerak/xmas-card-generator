from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import csv

INPUT="data.csv"
OUTPUT="xmas-cards.docx"

def load_csv(file_path):
    data = []

    try:
        with open(file_path, encoding='utf-8') as csv_file:
            csv_reader = csv.DictReader(csv_file)
            for row in csv_reader:
                data.append(row)
    except FileNotFoundError:
        print(f"Error: File not found at {file_path}")
    except Exception as e:
        print(f"An error occurred: {e}")

    return data

# static data for B type (vykani)
head_b = "V Pardubicích dne 10. prosince 2024\n\n"
body_b = (
    "Milí dobrovolníci, dobrodinci a pomocníci,\n\n"
    "dovolte mi bych vám poděkoval za vaši práci ve farnosti a zapojení se do života našeho společenství.  Jako poděkování přijměte, prosím, pozvání na mši svatou v pátek 17. ledna v 18:00 v kostele sv. Bartoloměje. Bude slavena za vás všechny, kteří jakýmkoliv způsobem aktivně pomáháte ve farnosti. Po bohoslužbě jste zváni na faru, kde u malého občerstvení bude naše společné setkání pokračovat.\n\n"
    "Svoji vděčnost bych vám chtěl vyjádřit i věnováním jednoho citátu z Písma. Ač je dopis psán pro všechny stejný, tento citát je pro každého osobní. Není vylosován, ale vybrán. Za každého z vás jsem se alespoň chvíli modlil, promítal jsem si v duchu váš život a přemýšlel jaká věta z Božího slova vyjadřuje vaši službu či životní situaci. Úmyslem citátu bylo vás potěšit, ocenit vaši službu, nebo dodat odvahu do nového roku. Pokud tomu tak není, a při jeho čtení pociťujete něco jiného, tak se omlouvám. K novoročence přikládám i dopis, který každý rok píši svým přátelům a známým.\n\n"
    "Děkuji vám za vaši službu a Bůh, který vidí i to co je skryté vám vše štědře odplatí!\n\n"
    "S vděčností Vám žehná\n\n"
    "P. Jan Uhlíř"
)

name_b = ", jako můj osobní dar přijměte toto Boží slovo:"

# static data for A type (tykani)
head_a = "V Pardubicích dne 10. prosince 2024\n\n"
body_a = (
    "Milí dobrovolníci, dobrodinci a pomocníci,\n\n"
    "dovolte mi bych vám poděkoval za vaši práci ve farnosti a zapojení se do života našeho společenství.  Jako poděkování přijměte, prosím, pozvání na mši svatou v pátek 17. ledna v 18:00 v kostele sv. Bartoloměje. Bude slavena za vás všechny, kteří jakýmkoliv způsobem aktivně pomáháte ve farnosti. Po bohoslužbě jste zváni na faru, kde u malého občerstvení bude naše společné setkání pokračovat. \n\n"
    "Svoji vděčnost bych vám chtěl vyjádřit i věnováním jednoho citátu z Písma. Ač je dopis psán pro všechny stejný, tento citát je pro každého osobní. Není vylosován, ale vybrán. Za každého z vás jsem se alespoň chvíli modlil, promítal jsem si v duchu váš život a přemýšlel jaká věta z Božího slova vyjadřuje vaši službu či životní situaci. Úmyslem citátu bylo vás potěšit, ocenit vaši službu, nebo dodat odvahu do nového roku. Pokud tomu tak není, a při jeho čtení pociťujete něco jiného, tak se omlouvám. K novoročence přikládám i dopis, který každý rok píši svým přátelům a známým.\n\n"
    "Děkuji vám za vaši službu a Bůh, který vidí i to co je skryté vám vše štědře odplatí!\n\n"
    "S vděčností Vám žehná \n\n"
    "P. Jan Uhlíř"
)

name_a = ", jako můj osobní dar přijmi toto Boží slovo:"



if __name__ == "__main__":

    # Load data from csv file
    data = load_csv(INPUT)

    # Initialize the document
    d = Document()

    # Set page size to A4
    section = d.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)

    # Set custom margins
    section.top_margin = Inches(0.79)
    section.left_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)
    section.right_margin = Inches(0.98)

    # Set default font style
    style = d.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(12)

    # Set line spacing to 1.5
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5


    for record in data:
        if(record["type"]=="A"):
            # Add greeting aligned to the right
            greeting = d.add_paragraph(head_a)
            greeting.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add main paragraph
            d.add_paragraph(body_a)

            # Add personalized section centered
            name_paragraph = d.add_paragraph(f'\n\n{record["address"]}{name_a}')
            name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add quote in italic and centered
            quote_paragraph = d.add_paragraph()
            quote_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            quote_run = quote_paragraph.add_run(f'„{record["quote"]}“')
            quote_run.italic = True

            # Add reference centered
            reference_paragraph = d.add_paragraph(record["source"])
            reference_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            # Add greeting aligned to the right
            greeting = d.add_paragraph(head_b)
            greeting.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add main paragraph
            d.add_paragraph(body_b)

            # Add personalized section centered
            name_paragraph = d.add_paragraph(f'\n\n{record["address"]}{name_b}')
            name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add quote in italic and centered
            quote_paragraph = d.add_paragraph()
            quote_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            quote_run = quote_paragraph.add_run(f'„{record["quote"]}“')
            quote_run.italic = True

            # Add reference centered
            reference_paragraph = d.add_paragraph(record["source"])
            reference_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        

        d.add_page_break()

    d.save(OUTPUT)